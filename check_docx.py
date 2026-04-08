import docx

def check_tags(filename):
    try:
        doc = docx.Document(filename)
        print(f"Checking {filename} for Jinja tags...")
        for i, p in enumerate(doc.paragraphs):
            if '{%' in p.text:
                print(f"Paragraph {i}: {p.text}")
        
        print("\nChecking tables...")
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, p in enumerate(cell.paragraphs):
                        if '{%' in p.text:
                            print(f"Table {t_idx}, Row {r_idx}, Cell {c_idx}, Paragraph {p_idx}: {p.text}")
    except Exception as e:
        print(f"Error reading {filename}: {e}")

if __name__ == '__main__':
    check_tags('plantilla_hv.docx')
